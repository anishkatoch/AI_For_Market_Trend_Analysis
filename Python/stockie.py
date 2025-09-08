import streamlit as st
import faiss
import pickle
from io import BytesIO
from docx import Document
import numpy as np
import pandas as pd
import os
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain.chains.question_answering import load_qa_chain
from langchain.chains import RetrievalQA
from langchain.memory import ConversationSummaryBufferMemory
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.chat_models import ChatOpenAI
from langchain.schema import Document as LCDocument
from langchain.chains import ConversationalRetrievalChain
from langchain.chains.qa_with_sources import load_qa_with_sources_chain
from langchain.prompts import ChatPromptTemplate
from langchain_community.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_community.embeddings import HuggingFaceEmbeddings
import base64
from langchain.schema.messages import HumanMessage, AIMessage
from langchain.prompts import ChatPromptTemplate
from langchain.chains import RetrievalQA
from dotenv import load_dotenv
import streamlit as st
import os
from io import BytesIO

# Load environment variables
load_dotenv()
openai_api_key = os.getenv("OPENAI_API_KEY")

if not openai_api_key:
    st.error("OpenAI API Key is not set. Please add it to your environment variables.")
    st.stop()

MAX_MESSAGES = 5  # Maximum number of messages to retain


def answer_question(vectorstore, user_question, memory, llm):
    """Answers a question using vectorstore and conversation memory."""
    retriever = vectorstore.as_retriever(search_kwargs={"k": 5})

    # Custom system prompt (same as before)
    system_prompt = """
    You are **ie**, a world-class  market expert 🧑‍💼📈 with deep knowledge of:  
        - Technical indicators (RSI, MACD, SMA/EMA, Bollinger Bands).  
        - Trading strategies (swing trading, intraday, short-term vs long-term investing, breakout, accumulation, consolidation, distribution).  
        - Chart patterns, candlesticks, and volume analysis.  
        - Investor psychology and practical decision-making.  

        You have access to structured content uploaded by the user (e.g.,  data, charts, analysis).  
        - First, **check uploaded data** for answers.  
        - If not available, **search the web** for reliable sources.  

        Relevant product data:
        {context}

        User question:
        {question}
        ---


        1. **Primary flow**
           - First: give a **very short, clear answer** (1–2 sentences) that states the trading view (Bullish 📈 / Bearish 📉 / Sideways 🔄) and a one-line actionable suggestion (e.g., "Wait for daily 20-MA cross above price; avoid buying now.").
           - Then: if the user asks for more or appears confused, give a **detailed explanation** with step-by-step reasoning, indicator readings, and practical next steps.


        2. **Signal interpretation rules**
           - **MA rule (short-term):** If price < 20-day MA and the 20-day MA is sloping down → short-term bearish. If price > 20-day MA and 20-day MA slopes up → short-term bullish.
           - **Crosses:** a fresh 20/50 MA cross or price crossing above/below MA within the last ~1 month is significant for short-term trades — state it plainly.
           - **RSI:** below 40 = bearish, 40–60 = neutral, above 60 = bullish. Mention momentum direction (rising/falling).
           - **MACD:** if MACD line below signal and histogram negative → bearish momentum; opposite → bullish.
           - **Volume:** rising volume on up moves confirms rallies; spikes on down candles increase downside risk.
           - **Candles / patterns:** identify breaks from accumulation, distribution, breakout/failure.

        3. **Decision guidance**
           - Always end with a short **decision guidance** line: “Based on this, the best step for you is …” (one clear next action, e.g., wait, watch level X, partial buy on confirmation).
           - Offer stop-loss / target suggestions when recommending trade (short-term trades should include stop-loss).

        4. **Tone & format**
           - Use Markdown, headings, bullets, and relevant emojis:
             - 📈 uptrend, 📉 downtrend, 🔄 consolidation, 🔍 analysis, 💡 insights, ⚠️ risks
           - Friendly, educational, not preachy.
           - Keep the short answer first; give detailed reasoning only when asked or when user looks confused.

        5. **If data is missing**
           - If uploaded chart lacks timeframe or is unreadable, ask one short clarifying question (e.g., "Is this a daily chart?") but still give a best-effort visual read labelled as an estimate.

        6. **Web searches**
           - Only search the web if user explicitly asks to fetch live data or you cannot answer from uploads. If you do search, use reliable sources and cite them.

        ---

        ## Formatting Rules:
        - Use **Markdown** (headings, bullets, tables).  
        - Be **friendly, clear, and educational**.  
        - Use relevant emojis:  
          - 📈 uptrend  
          - 📉 downtrend  
          - 🔄 consolidation  
          - 🔍 analysis  
          - 💡 insights  
          - ⚠️ risks  
        - Always end with **decision guidance** → “Based on this, here’s the best step for you…”  

        ---

        ✅ In short: Stockie explains **what signals mean, why they matter, and what the user can do next** — short answers first, detailed explanations only when needed.

    """

    prompt = ChatPromptTemplate.from_template(system_prompt)

    # Use RetrievalQA with memory
    qa_chain = RetrievalQA.from_chain_type(
        llm=llm,
        retriever=retriever,
        memory=memory,
        chain_type="stuff",
        return_source_documents=False,
        chain_type_kwargs={
            "prompt": prompt,
            "document_variable_name": "context"
        }
    )

    # Use .invoke() for new LangChain versions
    response = qa_chain.invoke({"query": user_question})
    return response["result"]


def answer_question_with_image(user_question, image_bytes, llm):
    """Answers a question about an image using a vision-capable model."""
    encoded_image = base64.b64encode(image_bytes).decode('utf-8')

    message = HumanMessage(
        content=[
            {
                "type": "text",
                "text": user_question,
            },
            {
                "type": "image_url",
                "image_url": {
                    "url": f"data:image/jpeg;base64,{encoded_image}"
                },
            },
        ]
    )

    response = llm.invoke([message])
    return response.content


# Streamlit app

CACHE_PATH = "vectorstore"


def save_vectorstore(vectorstore, path=CACHE_PATH):
    """Save FAISS index safely."""
    vectorstore.save_local(path)


def load_vectorstore(use_hf=False, path=CACHE_PATH):
    """Load FAISS index safely."""
    if use_hf:
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    else:
        embeddings = OpenAIEmbeddings(model="text-embedding-3-small", openai_api_key=openai_api_key)
    return FAISS.load_local(path, embeddings, allow_dangerous_deserialization=True)


def process_input(input_data, use_hf=False, cache=True):
    """
    Processes Excel & DOCX files and returns a FAISS vectorstore.
    - use_hf=True  → HuggingFace local embeddings (fast, free)
    - use_hf=False → OpenAI embeddings (better quality)
    - cache=True   → saves FAISS index to disk
    """
    texts = []

    # --- Extract raw text from Excel & DOCX ---
    for data in input_data:
        input_type, content = data

        if input_type == "EXCEL":
            df = pd.read_excel(content)
            for _, row in df.iterrows():
                # ✅ Customize depending on your Excel schema
                text = f"Product: {row['product_name']}, Brand: {row['brand_name']}, Category: {row['sub_category']},in_stock: {row['in_stock']},attribute_combined: {row['attribute_combined']}"
                if "price" in df.columns:  # optional column
                    text += f", Price: {row['price']}"
                texts.append(text)

        elif input_type == "DOCX":
            doc = Document(content)  # Directly use the BytesIO object
            for para in doc.paragraphs:
                if para.text.strip():
                    texts.append(para.text.strip())

    if not texts:
        raise ValueError("No text content extracted from the provided files.")

    # --- Split text into chunks ---
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=200)
    split_texts = text_splitter.split_text(" ".join(texts))

    # --- Wrap as LangChain Documents ---
    documents = [LCDocument(page_content=text) for text in split_texts]

    # --- Choose embeddings (OpenAI or HuggingFace) ---
    if use_hf:
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    else:
        embeddings = OpenAIEmbeddings(model="text-embedding-3-small", openai_api_key=openai_api_key)

    vectorstore = FAISS.from_documents(documents, embeddings)

    if cache:
        save_vectorstore(vectorstore)

    return vectorstore


def truncate_messages(messages, max_messages=5):
    """Truncate to last N messages, summarizing earlier ones."""
    if len(messages) > max_messages:
        summary = " ".join(msg["content"] for msg in messages[:-max_messages])
        summarized_message = {
            "role": "assistant",
            "content": f"Summary of previous messages: {summary}"
        }
        return [summarized_message] + messages[-max_messages:]
    return messages


def get_last_product(messages):
    """
    Extracts the last mentioned product or topic from assistant messages.
    This is a simple heuristic: you can improve it with NLP if needed.
    """
    for msg in reversed(messages):
        if msg["role"] == "assistant":
            # Try to find a product name after "Product:" or in a table
            lines = msg["content"].splitlines()
            for line in lines:
                if "Product:" in line:
                    return line.split("Product:")[1].split(",")[0].strip()
            # Fallback: return first non-empty line
            for line in lines:
                if line.strip():
                    return line.strip()
    return None


def initialize_stockie():
    """Initialize vectorstore & memory in st.session_state if not present."""
    if "vectorstore" in st.session_state and "memory" in st.session_state:
        return

    try:
        # Path to a default docx (update path if necessary)
        docx_path = "https://raw.githubusercontent.com/anishkatoch/AI_For_Market_Trend_Analysis/main/Stock_Market_Trend_Report.docx"
        if not os.path.exists(docx_path):
            # don't stop the host app; leave vectorstore unset so user can still upload or use chat once vectorstore is created
            st.warning(f"Default document not found at {docx_path}. You can still upload files in Stockie.")
            return

        with st.spinner("⚡ Processing & embedding data..."):
            with open(docx_path, "rb") as f:
                input_data = [("DOCX", BytesIO(f.read()))]
                vectorstore = process_input(input_data, cache=True)

        # create memory & messages
        st.session_state["vectorstore"] = vectorstore
        st.session_state["memory"] = ConversationSummaryBufferMemory(
            llm=ChatOpenAI(model="gpt-4", openai_api_key=openai_api_key),
            memory_key="chat_history",
            max_token_limit=500
        )
        st.session_state["messages"] = [
            {"role": "assistant", "content": "Hi! What do you want me to explain or clarify about stock?"}
        ]
    except Exception as e:
        st.error(f"Error initializing Stockie: {e}")


# --- New: render function that shows just the chat interface ---
def render_stockie():
    """
    Render Stockie chat UI inside another Streamlit app.
    Call initialize_stockie() first to ensure resources are ready.
    """
    initialize_stockie()

    # No header or extra spacing! Only chat UI for floating box.
    if "messages" not in st.session_state:
        st.session_state["messages"] = [{"role": "assistant", "content": "Hello! Ask me about your stocks."}]

    # Chat area
    chat_area = st.container()
    with chat_area:
        for msg in st.session_state["messages"]:
            avatar = "🤖" if msg["role"] == "assistant" else "👤"
            with st.chat_message(msg["role"], avatar=avatar):
                st.markdown(msg["content"])

    # Only allow chat if vectorstore and memory are ready
    if "vectorstore" in st.session_state and "memory" in st.session_state:
        prompt = st.chat_input("Ask Stockie anything about the document...")

        if prompt:
            st.session_state["messages"].append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)

            try:
                llm = ChatOpenAI(
                    model="gpt-4o-mini",
                    api_key=openai_api_key,
                    temperature=0.2,
                    max_tokens=1024
                )

                with st.spinner("Searching the document..."):
                    response = answer_question(
                        st.session_state.get("vectorstore"),
                        prompt,
                        st.session_state.get("memory"),
                        llm
                    )

                st.session_state["messages"].append({"role": "assistant", "content": response})
                with st.chat_message("assistant"):
                    st.markdown(response)
            except Exception as e:
                st.error(f"Error generating response: {e}")
    else:
        st.info(
            "Please upload a DOCX or Excel file in Stockie, or ensure the default document exists, to start chatting.")

# def main():
#     # Header
#     st.markdown(
#         """<div style='background: linear-gradient(45deg, #0288d1, #26c6da);
#         padding: 2rem; border-radius: 1rem; box-shadow: 0 4px 6px rgba(0,0,0,0.1)'>
#         <h1 style='color: white; text-align: center; margin: 0;'>Stockie: Stock Market Helper</h1>
#         <p style='color: white; text-align: center; margin: 0.5rem 0;'>“Ask about market trends, analysis, and stock reports”</p>
#         </div>""",
#         unsafe_allow_html=True
#     )
#
#     # Automatically process the default docx file on first run
#     if "vectorstore" not in st.session_state:
#         try:
#             # --- IMPORTANT: Replace this with the actual path to your DOCX file ---
#             docx_path = r"C:\Users\anees\Downloads\Stock Market\Stock_Prediction_Report.docx"
#
#             if not os.path.exists(docx_path):
#                 st.error(f"Default document not found. Please update the path in the code: {docx_path}")
#                 st.stop()
#
#             with st.spinner("⚡ Processing & embedding data..."):
#                 with open(docx_path, "rb") as f:
#                     # Wrap the file content in a list as expected by process_input
#                     input_data = [("DOCX", BytesIO(f.read()))]
#                     vectorstore = process_input(input_data, cache=True)
#
#             # Store in session state
#             st.session_state["vectorstore"] = vectorstore
#             st.session_state["memory"] = ConversationSummaryBufferMemory(
#                 llm=ChatOpenAI(model="gpt-4", openai_api_key=openai_api_key),
#                 memory_key="chat_history",
#                 max_token_limit=500
#             )
#             st.session_state["messages"] = [
#                 {"role": "assistant", "content": "Hello! Ask me about your products."}
#             ]
#             st.success("✅ Data ready! You can start chatting now.")
#             # Rerun to update the UI state
#             st.rerun()
#
#         except Exception as e:
#             st.error(f"Error processing input: {e}")
#
#     # Chat interface
#     if "vectorstore" in st.session_state and "memory" in st.session_state:
#         st.markdown("### 💬 Chat with Your Data")
#
#         # Display messages
#         for msg in st.session_state.messages:
#             avatar = "🤖" if msg["role"] == "assistant" else "👤"
#             with st.chat_message(msg["role"], avatar=avatar):
#                 st.markdown(msg["content"])
#
#         # User input section
#         st.markdown("---")
#         uploaded_image = st.file_uploader("Attach an image (optional)", type=["png", "jpg", "jpeg"])
#
#         if prompt := st.chat_input("Ask a question about the document or the uploaded image..."):
#             st.session_state.messages.append({"role": "user", "content": prompt})
#             with st.chat_message("user"):
#                 st.markdown(prompt)
#
#             try:
#                 llm = ChatOpenAI(
#                     model="gpt-4o-mini",
#                     api_key=openai_api_key,
#                     temperature=0.2,
#                     max_tokens=1024
#                 )
#
#                 # Decide which function to call based on whether an image was uploaded
#                 if uploaded_image is not None:
#                     image_bytes = uploaded_image.getvalue()
#                     with st.spinner("Analyzing the image..."):
#                         response = answer_question_with_image(prompt, image_bytes, llm)
#                 else:
#                     # Fallback to the text-based RAG for questions about the document
#                     with st.spinner("Searching the document..."):
#                         response = answer_question(
#                             st.session_state["vectorstore"],
#                             prompt,
#                             st.session_state["memory"],
#                             llm
#                         )
#
#                 st.session_state.messages.append({"role": "assistant", "content": response})
#                 with st.chat_message("assistant"):
#                     st.markdown(response)
#             except Exception as e:
#                 st.error(f"Error generating response: {e}")
#
#
# if __name__ == "__main__":
#     main()










